"""
Resume parsing: extract raw text + structured sections from PDF/DOCX uploads.

Two-stage output:
  1. raw_text  — full resume as a single string (for embedding-based extraction)
  2. sections  — dict of detected section text keyed by section name
                 (for section-weighted skill extraction: skills section = 3×)
"""

import io
import re
from pathlib import Path
from typing import Optional

import structlog

log = structlog.get_logger(__name__)

# Section header patterns — order matters (most specific first)
_SECTION_PATTERNS = [
    ("skills", r"\b(technical\s+)?skills?\b|\bcore\s+competenc"),
    ("experience", r"\bwork\s+experience\b|\bprofessional\s+experience\b|\bemployment\b|\bexperience\b"),
    ("education", r"\beducation\b|\bacademic\b|\bdegree\b"),
    ("projects", r"\bprojects?\b|\bpersonal\s+projects?\b|\bside\s+projects?\b"),
    ("certifications", r"\bcertif"),
    ("summary", r"\bsummary\b|\bobjective\b|\bprofile\b|\babout\b"),
    ("publications", r"\bpublications?\b|\bresearch\b|\bpapers?\b"),
    ("awards", r"\bawards?\b|\bachievements?\b|\bhonors?\b"),
]
_SECTION_RE = [(name, re.compile(pat, re.IGNORECASE)) for name, pat in _SECTION_PATTERNS]

# Maximum text length we'll process (prevent abuse / OOM on 100-page PDFs)
_MAX_TEXT_CHARS = 50_000


class ParseError(Exception):
    pass


def parse_resume(file_bytes: bytes, filename: str) -> dict:
    """
    Parse a resume file (PDF or DOCX) and return:
        {
            "raw_text": str,
            "sections": {"skills": str, "experience": str, ...},
            "layout_type": "pdf" | "docx",
            "char_count": int,
        }

    Raises ParseError on unsupported format or extraction failure.
    """
    suffix = Path(filename).suffix.lower()

    if suffix == ".pdf":
        raw_text = _extract_pdf(file_bytes)
        layout_type = "pdf"
    elif suffix in (".docx", ".doc"):
        raw_text = _extract_docx(file_bytes)
        layout_type = "docx"
    else:
        raise ParseError(f"Unsupported file type: {suffix!r}. Upload a PDF or DOCX.")

    if not raw_text or not raw_text.strip():
        raise ParseError("No text could be extracted from the file. The resume may be image-based or corrupted.")

    raw_text = raw_text[:_MAX_TEXT_CHARS]
    sections = _detect_sections(raw_text)

    log.info(
        "resume_parsed",
        filename=filename,
        layout_type=layout_type,
        char_count=len(raw_text),
        sections_found=list(sections.keys()),
    )

    return {
        "raw_text": raw_text,
        "sections": sections,
        "layout_type": layout_type,
        "char_count": len(raw_text),
    }


# ---------------------------------------------------------------------------
# Format extractors
# ---------------------------------------------------------------------------

def _extract_pdf(file_bytes: bytes) -> str:
    try:
        import fitz  # PyMuPDF

        doc = fitz.open(stream=file_bytes, filetype="pdf")
        pages = []
        for page in doc:
            pages.append(page.get_text("text"))
        doc.close()
        return "\n".join(pages)
    except ImportError:
        raise ParseError("PDF parsing library (PyMuPDF) not installed.")
    except Exception as exc:
        raise ParseError(f"Failed to parse PDF: {exc}") from exc


def _extract_docx(file_bytes: bytes) -> str:
    try:
        from docx import Document

        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also capture table cells (skills often live in tables)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return "\n".join(paragraphs)
    except ImportError:
        raise ParseError("DOCX parsing library (python-docx) not installed.")
    except Exception as exc:
        raise ParseError(f"Failed to parse DOCX: {exc}") from exc


# ---------------------------------------------------------------------------
# Section detector
# ---------------------------------------------------------------------------

def _detect_sections(text: str) -> dict[str, str]:
    """
    Split resume text into named sections by detecting section headers.
    Returns a dict mapping section name → section text.
    Unmatched text goes into 'other'.
    """
    lines = text.splitlines()
    # Find (line_index, section_name) for each detected header
    boundaries: list[tuple[int, str]] = []
    for i, line in enumerate(lines):
        stripped = line.strip()
        if not stripped or len(stripped) > 80:
            continue
        for name, pattern in _SECTION_RE:
            if pattern.search(stripped):
                # Avoid false positives: header lines are short and often ALL-CAPS or title-case
                is_header = (
                    stripped.isupper()
                    or stripped.istitle()
                    or re.match(r"^[A-Z][a-zA-Z\s&/]+$", stripped)
                    or len(stripped) < 40
                )
                if is_header:
                    boundaries.append((i, name))
                    break

    if not boundaries:
        return {"other": text}

    sections: dict[str, str] = {}
    # Text before the first header
    if boundaries[0][0] > 0:
        preamble = "\n".join(lines[: boundaries[0][0]]).strip()
        if preamble:
            sections["summary"] = preamble

    for idx, (line_no, section_name) in enumerate(boundaries):
        end = boundaries[idx + 1][0] if idx + 1 < len(boundaries) else len(lines)
        body = "\n".join(lines[line_no + 1 : end]).strip()
        if body:
            # If the same section appears twice (e.g. two "Skills" headings), concatenate
            if section_name in sections:
                sections[section_name] += "\n" + body
            else:
                sections[section_name] = body

    return sections


def get_section_weight(section_name: Optional[str]) -> float:
    """Return the extraction weight for a given section (skills = 3×, others = 1×)."""
    weights = {
        "skills": 3.0,
        "projects": 1.5,
        "experience": 1.0,
        "education": 0.5,
        "certifications": 1.2,
        "summary": 0.8,
        "other": 0.7,
    }
    return weights.get(section_name or "other", 1.0)
