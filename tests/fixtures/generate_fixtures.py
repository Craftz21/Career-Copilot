#!/usr/bin/env python3
"""
Generate synthetic resume fixtures for parser evaluation.

Creates programmatic PDF/DOCX files that represent each resume class.
Each fixture is designed to match its paired ground_truth/*.json file exactly.

Run once before executing the evaluator:
    python tests/fixtures/generate_fixtures.py

Requires: PyMuPDF (fitz), python-docx
"""

import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent.parent
sys.path.insert(0, str(ROOT))

TEST_RESUMES = ROOT / "test_resumes"

_PDF_CLASSES = ["single_column", "double_column", "ats", "canva", "image_heavy"]
_DOCX_CLASSES = ["docx"]


def _ensure_dirs():
    for cls in _PDF_CLASSES + _DOCX_CLASSES:
        (TEST_RESUMES / cls).mkdir(parents=True, exist_ok=True)


# ─────────────────── PyMuPDF helpers ─────────────────────────────────────────

def _pdf_single_col(path: Path, author: str, contact: str, sections: list[tuple[str, str]]):
    """
    Standard single-column PDF. Uppercase section headers, body indented.
    This is the baseline — parser should score near-perfectly on this layout.
    """
    import fitz

    doc = fitz.open()
    page = doc.new_page(width=612, height=792)

    y = 55
    page.insert_text((50, y), author, fontsize=14, fontname="helv")
    y += 20
    page.insert_text((50, y), contact, fontsize=9)
    y += 28

    for header, body in sections:
        page.insert_text((50, y), header.upper(), fontsize=11, fontname="helv")
        y += 16

        for raw_line in body.split("\n"):
            line = raw_line.strip()
            if not line:
                y += 7
                continue
            # Naive word-wrap at 88 chars
            while len(line) > 88:
                cut = line[:88].rfind(" ")
                cut = cut if cut != -1 else 88
                page.insert_text((62, y), line[:cut], fontsize=10)
                y += 13
                line = line[cut:].strip()
            page.insert_text((62, y), line, fontsize=10)
            y += 13
        y += 8

    doc.save(str(path))
    doc.close()


def _pdf_double_col(path: Path, left_lines: list[str], right_lines: list[str]):
    """
    Two-column layout where both columns share overlapping Y ranges.
    PyMuPDF's get_text('text') reads by Y then X, interleaving the columns.
    Headers use '##' prefix to signal bold/uppercase rendering.
    """
    import fitz

    doc = fitz.open()
    page = doc.new_page(width=612, height=792)

    y = 52
    for line in left_lines:
        if line.startswith("##"):
            page.insert_text((50, y), line[2:].strip().upper(), fontsize=11, fontname="helv")
            y += 17
        elif line == "":
            y += 8
        else:
            page.insert_text((56, y), line, fontsize=10)
            y += 13

    y = 52
    for line in right_lines:
        if line.startswith("##"):
            page.insert_text((310, y), line[2:].strip().upper(), fontsize=11, fontname="helv")
            y += 17
        elif line == "":
            y += 8
        else:
            page.insert_text((316, y), line, fontsize=10)
            y += 13

    doc.save(str(path))
    doc.close()


def _pdf_image_only(path: Path, readable_text: str):
    """
    PDF where all content is a rasterized bitmap image — no text layer.
    parse_resume() MUST raise ParseError on this (no OCR available).
    """
    import fitz

    # Render text → pixmap
    tmp = fitz.open()
    tmp_page = tmp.new_page(width=612, height=792)
    y = 80
    for line in readable_text.split("\n"):
        tmp_page.insert_text((50, y), line.strip(), fontsize=10)
        y += 14
    pix = tmp_page.get_pixmap(matrix=fitz.Matrix(1.5, 1.5))
    tmp.close()

    # Embed pixmap as image — no selectable text
    doc = fitz.open()
    pg = doc.new_page(width=612, height=792)
    pg.insert_image(fitz.Rect(0, 0, 612, 792), pixmap=pix)
    doc.save(str(path))
    doc.close()


# ─────────────────── python-docx helpers ─────────────────────────────────────

def _docx_with_table(
    path: Path,
    author: str,
    skills_table: list[tuple[str, str]],
    sections: list[tuple[str, str]],
):
    """DOCX with a two-column skills table (Category | Skills rows)."""
    from docx import Document

    doc = Document()
    doc.add_heading(author, level=1)
    doc.add_paragraph("contact@example.com | linkedin.com/in/profile")

    doc.add_heading("Skills", level=2)
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    hdr[0].text = "Category"
    hdr[1].text = "Technologies"
    for cat, techs in skills_table:
        row = tbl.add_row().cells
        row[0].text = cat
        row[1].text = techs

    doc.add_paragraph("")
    for header, body in sections:
        doc.add_heading(header, level=2)
        for line in body.split("\n"):
            if line.strip():
                doc.add_paragraph(line.strip())

    doc.save(str(path))


def _docx_plain(path: Path, author: str, sections: list[tuple[str, str]]):
    """Plain DOCX — Word Heading styles, no tables."""
    from docx import Document

    doc = Document()
    doc.add_heading(author, level=1)
    doc.add_paragraph("contact@email.com")

    for header, body in sections:
        doc.add_heading(header, level=2)
        for line in body.split("\n"):
            if line.strip():
                doc.add_paragraph(line.strip())

    doc.save(str(path))


# ─────────────────── Fixture definitions ─────────────────────────────────────
# Each fixture must exactly match its paired ground_truth/*.json file.

def gen_sc_01():
    """sc_01_standard — single column, easy."""
    _pdf_single_col(
        TEST_RESUMES / "single_column" / "sc_01_standard.pdf",
        author="Jane Smith",
        contact="jane.smith@example.com | linkedin.com/in/janesmith",
        sections=[
            ("Summary",
             "Backend software engineer with 4 years building scalable REST APIs."),
            ("Skills",
             "Python, FastAPI, PostgreSQL, Docker, Redis, Git, Linux"),
            ("Experience",
             "Senior Backend Engineer | TechCorp Inc | 2022 - Present\n"
             "- Developed REST APIs using FastAPI and Python\n"
             "- Optimized PostgreSQL queries, reducing p99 latency by 40%\n"
             "- Containerized all services using Docker\n"
             "\n"
             "Software Engineer | StartupXYZ | 2020 - 2022\n"
             "- Built Python microservices with Redis caching layer"),
            ("Education",
             "B.S. Computer Science | Stanford University | 2020"),
            ("Projects",
             "DataPipeline - ETL pipeline using Python and PostgreSQL\n"
             "APIGateway - Rate-limited gateway with Redis"),
        ],
    )


def gen_sc_02():
    """sc_02_dense — single column, medium (non-standard header names)."""
    _pdf_single_col(
        TEST_RESUMES / "single_column" / "sc_02_dense.pdf",
        author="Alex Chen",
        contact="alex.chen@example.com | github.com/alexchen",
        sections=[
            ("About Me",
             "Full-stack developer specializing in React and Node.js ecosystems."),
            ("Technical Skills",
             "JavaScript, TypeScript, React, Node.js, MongoDB, AWS, Docker, Git"),
            ("Work Experience",
             "Frontend Developer | WebAgency | 2021 - Present\n"
             "- Built responsive web applications using React and TypeScript\n"
             "- Integrated with Node.js backend REST APIs\n"
             "\n"
             "Junior Developer | FreelanceCo | 2019 - 2021\n"
             "- Developed MongoDB schemas and Node.js REST APIs"),
            ("Education Background",
             "B.S. Information Technology | MIT | 2021"),
            ("Side Projects",
             "E-commerce site: React, Node.js, MongoDB\n"
             "Portfolio: deployed on AWS S3 + CloudFront"),
        ],
    )


def gen_dc_01():
    """dc_01_sidebar — double column, hard."""
    left = [
        "## Skills",
        "Python",
        "Docker",
        "Kubernetes",
        "Redis",
        "PostgreSQL",
        "",
        "## Education",
        "B.S. Computer Science",
        "UC Berkeley, 2021",
    ]
    right = [
        "## Michael Brown",
        "mike@example.com",
        "",
        "## Experience",
        "DevOps Engineer | CloudCo | 2021-Present",
        "- Orchestrated containers with Kubernetes",
        "- Deployed CI/CD pipelines using Docker",
        "- Managed PostgreSQL cluster across 3 regions",
        "- Deployed Redis caching layer for API responses",
        "",
        "## Projects",
        "K8s monitoring dashboard written in Python",
    ]
    _pdf_double_col(TEST_RESUMES / "double_column" / "dc_01_sidebar.pdf", left, right)


def gen_ats_01():
    """ats_01_clean — ATS safe, easy."""
    _pdf_single_col(
        TEST_RESUMES / "ats" / "ats_01_clean.pdf",
        author="Michael Johnson",
        contact="michael.j@example.com | (555) 000-1234",
        sections=[
            ("Technical Skills",
             "Languages: Python, Go, JavaScript\n"
             "Frameworks: FastAPI, React, Node.js\n"
             "Databases: PostgreSQL, Redis, MongoDB\n"
             "DevOps: Docker, Kubernetes, Git"),
            ("Work Experience",
             "Backend Engineer | BigTech Corp | 2020 - Present\n"
             "- Designed distributed systems using Python and Go\n"
             "- Maintained PostgreSQL and Redis infrastructure\n"
             "\n"
             "Software Engineer | MidTech | 2018 - 2020\n"
             "- Built React frontend and Node.js backend APIs"),
            ("Education",
             "M.S. Computer Science | UC Berkeley | 2020\n"
             "B.S. Computer Science | UCLA | 2018"),
            ("Certifications",
             "AWS Certified Solutions Architect | 2023"),
        ],
    )


def gen_ats_02():
    """ats_02_minimal — ATS minimal, easy."""
    _pdf_single_col(
        TEST_RESUMES / "ats" / "ats_02_minimal.pdf",
        author="Sarah Lee",
        contact="sarah.lee@example.com",
        sections=[
            ("Skills",
             "Python, SQL, Tableau, Pandas, NumPy, Machine Learning, Scikit-Learn"),
            ("Experience",
             "Data Analyst | DataCorp | 2022 - Present\n"
             "- Analyzed large datasets using Python and Pandas\n"
             "- Built Tableau dashboards for executive reporting\n"
             "- Applied machine learning models with Scikit-Learn"),
            ("Education",
             "B.S. Statistics | NYU | 2022"),
        ],
    )


def gen_canva_01():
    """canva_01_creative — Canva style with sidebar, decorative header, hard."""
    import fitz

    path = TEST_RESUMES / "canva" / "canva_01_creative.pdf"
    doc = fitz.open()
    page = doc.new_page(width=612, height=792)

    # Decorative header band
    page.draw_rect(fitz.Rect(0, 0, 612, 80), color=None, fill=(0.18, 0.38, 0.78))
    page.insert_text((28, 42), "PRIYA SHARMA", fontsize=18, color=(1, 1, 1), fontname="helv")
    page.insert_text((28, 62), "Full Stack Developer", fontsize=11, color=(0.88, 0.92, 1.0))

    # Contact scattered (Canva style — no fixed line order)
    page.insert_text((370, 38), "priya@example.com", fontsize=9, color=(1, 1, 1))
    page.insert_text((370, 52), "github.com/priyasharma", fontsize=9, color=(1, 1, 1))

    # Left sidebar
    page.draw_rect(fitz.Rect(0, 80, 185, 792), color=None, fill=(0.94, 0.95, 0.98))

    page.insert_text((14, 108), "SKILLS", fontsize=10, fontname="helv")
    sidebar_skills = ["Python", "React", "MongoDB", "Docker", "FastAPI", "Git", "JavaScript"]
    sy = 126
    for s in sidebar_skills:
        page.insert_text((14, sy), f"• {s}", fontsize=9)
        sy += 14

    page.insert_text((14, sy + 10), "EDUCATION", fontsize=10, fontname="helv")
    sy += 28
    page.insert_text((14, sy), "B.S. Computer Science", fontsize=9)
    sy += 13
    page.insert_text((14, sy), "IIT Delhi, 2023", fontsize=9)

    # Main content right of sidebar
    page.insert_text((200, 108), "EXPERIENCE", fontsize=11, fontname="helv")
    exp = [
        "Full Stack Developer | StartupAI | 2023-Present",
        "  - Built Python backend APIs with FastAPI",
        "  - Developed React frontend with TypeScript",
        "  - Managed MongoDB collections and indexes",
        "  - Containerized with Docker Compose",
        "",
        "Intern Developer | WebCo | 2022",
        "  - JavaScript and React component development",
    ]
    ey = 126
    for line in exp:
        page.insert_text((200, ey), line, fontsize=9)
        ey += 13

    page.insert_text((200, ey + 10), "PROJECTS", fontsize=11, fontname="helv")
    ey += 28
    page.insert_text((200, ey), "AIChat — Python, FastAPI, MongoDB, React", fontsize=9)
    ey += 14
    page.insert_text((200, ey), "Portfolio — React, Docker, deployed on AWS", fontsize=9)

    doc.save(str(path))
    doc.close()


def gen_docx_01():
    """docx_01_tables — DOCX with a skills table, medium."""
    _docx_with_table(
        TEST_RESUMES / "docx" / "docx_01_tables.docx",
        author="David Kim",
        skills_table=[
            ("Languages",   "Python, JavaScript, TypeScript, SQL"),
            ("Frameworks",  "FastAPI, React, Node.js, Express"),
            ("Databases",   "PostgreSQL, MongoDB, Redis"),
            ("DevOps",      "Docker, Git, Linux, AWS"),
        ],
        sections=[
            ("Experience",
             "Software Engineer | Fintech Inc | 2021-Present\n"
             "- Built financial APIs using FastAPI and Python\n"
             "- Designed PostgreSQL schemas for transaction data\n"
             "- Implemented Redis caching, cutting response time by 60%\n"
             "\n"
             "Junior Developer | Agency | 2019-2021\n"
             "- React and Node.js full-stack development"),
            ("Education",
             "B.S. Computer Engineering | Caltech | 2019"),
            ("Projects",
             "BudgetTracker: Python, FastAPI, PostgreSQL, React frontend"),
        ],
    )


def gen_docx_02():
    """docx_02_plain — plain DOCX, easy."""
    _docx_plain(
        TEST_RESUMES / "docx" / "docx_02_plain.docx",
        author="Emily Zhang",
        sections=[
            ("Summary",
             "Machine learning engineer with expertise in Python and deep learning frameworks."),
            ("Skills",
             "Python, TensorFlow, PyTorch, Scikit-Learn, Pandas, NumPy, SQL, Docker, Git"),
            ("Experience",
             "ML Engineer | AIStartup | 2022-Present\n"
             "Trained neural networks using TensorFlow and PyTorch.\n"
             "Feature engineering with Pandas and NumPy.\n"
             "Containerized model serving with Docker.\n"
             "\n"
             "Data Scientist | DataCo | 2020-2022\n"
             "Applied Scikit-Learn for classification and regression tasks."),
            ("Education",
             "M.S. Machine Learning | Carnegie Mellon University | 2020"),
        ],
    )


def gen_img_01():
    """img_01_scanned — image-only PDF, hard (expected ParseError)."""
    content = (
        "RESUME - SCANNED COPY\n"
        "John Scanned\n"
        "SKILLS: Python, Docker, PostgreSQL\n"
        "EXPERIENCE: Software Engineer 2020-2024\n"
        "  Built REST APIs and managed databases\n"
        "EDUCATION: B.S. Computer Science, 2020"
    )
    _pdf_image_only(TEST_RESUMES / "image_heavy" / "img_01_scanned.pdf", content)


# ─────────────────── Entry point ──────────────────────────────────────────────

def main():
    print("Creating directories...")
    _ensure_dirs()

    fitz_ok = True
    docx_ok = True

    try:
        import fitz  # noqa: F401
    except ImportError:
        print("[WARN] PyMuPDF not installed — skipping all PDF fixtures")
        fitz_ok = False

    try:
        from docx import Document  # noqa: F401
    except ImportError:
        print("[WARN] python-docx not installed — skipping DOCX fixtures")
        docx_ok = False

    generators = [
        (fitz_ok,  gen_sc_01,    "single_column/sc_01_standard.pdf"),
        (fitz_ok,  gen_sc_02,    "single_column/sc_02_dense.pdf"),
        (fitz_ok,  gen_dc_01,    "double_column/dc_01_sidebar.pdf"),
        (fitz_ok,  gen_ats_01,   "ats/ats_01_clean.pdf"),
        (fitz_ok,  gen_ats_02,   "ats/ats_02_minimal.pdf"),
        (fitz_ok,  gen_canva_01, "canva/canva_01_creative.pdf"),
        (fitz_ok,  gen_img_01,   "image_heavy/img_01_scanned.pdf"),
        (docx_ok,  gen_docx_01,  "docx/docx_01_tables.docx"),
        (docx_ok,  gen_docx_02,  "docx/docx_02_plain.docx"),
    ]

    ok = 0
    for enabled, fn, label in generators:
        if not enabled:
            print(f"  SKIP  {label}")
            continue
        try:
            fn()
            print(f"  OK    {label}")
            ok += 1
        except Exception as exc:
            print(f"  FAIL  {label}: {exc}")

    print(f"\n{ok}/{len(generators)} fixtures generated.")
    print("\nRun the evaluator:")
    print("  python scripts/evaluate_parser.py --resume-dir test_resumes/ --verbose")


if __name__ == "__main__":
    main()
